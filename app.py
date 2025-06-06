# -*- coding: utf-8 -*-
import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO

client = OpenAI(api_key=st.secrets["openai_api_key"])

st.title("🎓 Asistente Pedagógico IA")
st.subheader("Sesiones y Rúbricas personalizadas con inteligencia artificial")

# Entradas del docente
docente = st.text_input("👨‍🏫 Nombre del docente", "Jimmy Albino Meneses")
colegio = st.text_input("🏫 Nombre del colegio", "I.E. Dora Mayer")
grado = st.selectbox("📚 Grado", ["1°", "2°", "3°", "4°", "5°"])
competencia = st.selectbox("🧠 Competencia", [
    "Indaga mediante métodos científicos para construir conocimientos.",
    "Explica el mundo físico basándose en conocimientos sobre los seres vivos, materia y energía, biodiversidad, Tierra y universo.",
    "Diseña y construye soluciones tecnológicas para resolver problemas de su entorno."
])
titulo = st.text_input("📝 Título de la sesión", "¿Cómo podemos demostrar que el aire ocupa espacio?")

# Botón 1: Generar sesión
if st.button("📄 Generar sesión IA"):
    with st.spinner("Generando sesión con IA..."):

        prompt_sesion = f"""
Eres un especialista en educación secundaria en Perú. Genera una sesión de aprendizaje robusta para el área de Ciencia y Tecnología en {grado} de secundaria. Incluye:

1. Datos generales (Docente: {docente}, Colegio: {colegio}, Área: Ciencia y Tecnología, Grado: {grado})
2. Título de la sesión: {titulo}
3. Propósito
4. Competencia: {competencia}
5. Capacidades y desempeños sugeridos
6. Actividades divididas en: Inicio (15 min), Desarrollo (50 min), Cierre (25 min)
7. Estrategias y recursos
8. Evaluación formativa

Escribe de forma clara y estructurada para copiarlo en un documento Word.
"""

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt_sesion}],
            temperature=0.7,
            max_tokens=1500
        )

        resultado = response.choices[0].message.content

        doc = Document()
        doc.add_heading('SESIÓN DE APRENDIZAJE', 0)
        for linea in resultado.split("\n"):
            doc.add_paragraph(linea)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Sesión generada con éxito.")
        st.download_button("📥 Descargar sesión Word", data=buffer, file_name="sesion_IA_dinamica.docx")

# Botón 2: Generar rúbrica
def generar_rubrica():
    prompt_rubrica = f"""
Eres un especialista en evaluación educativa del área Ciencia y Tecnología en Perú. Crea una rúbrica de evaluación para estudiantes de {grado} de secundaria, alineada a la competencia "{competencia}", sobre el tema "{titulo}". Usa criterios claros, niveles de logro (AD, A, B, C) y descripciones observables para cada nivel. Formato tabla simple lista para exportar a Word.
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt_rubrica}],
        temperature=0.7,
        max_tokens=1000
    )

    resultado = response.choices[0].message.content

    doc = Document()
    doc.add_heading('RÚBRICA DE EVALUACIÓN', 0)
    for linea in resultado.split("\n"):
        doc.add_paragraph(linea)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Rúbrica generada con éxito.")
    st.download_button("📥 Descargar rúbrica Word", data=buffer, file_name="rubrica_IA.docx")

if st.button("📊 Generar rúbrica de evaluación"):
    with st.spinner("Generando rúbrica personalizada..."):
        generar_rubrica()
